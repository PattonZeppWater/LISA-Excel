"""
Populates SubmittalCoverLetter_Template_20260504.docx with form data.

Strategy:
  1. Load template bytes from disk
  2. Open with python-docx (in memory only — doc.save() is never called)
  3. Set SDT content controls (Submittal No. / date picker / Subject)
  4. Toggle YES/NO checkboxes for Response Requested
  5. Fill comment log tables (tables 2-6); remove unused ones entirely
  6. Rebuild markup summary subsection paragraphs
  7. Serialize doc.element to bytes (document.xml only — no PackageWriter)
  8. Reconstruct the output ZIP from the template, copying every part verbatim:
       - word/document.xml     → swapped for the python-docx result
       - docProps/custom.xml   → property values written with actual field data
                                 (headers/footers display via DOCPROPERTY field codes)
       - word/settings.xml     → w:updateFields added so Word refreshes fields on open
       - [Content_Types].xml   → lisa_payload.json Override entry inserted
       - word/lisa_payload.json→ added with full form payload
       - everything else       → copied verbatim (customXml/, glossary/, rels/, media…)

This preserves all template parts that python-docx does not model so Word never
triggers the "content not recoverable" repair dialog.
"""

import copy
import html
import io
import json
import os
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as _Table
from lxml import etree

TEMPLATE_PATH = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "..", "_Templates",
                 "SubmittalCoverLetter_Template_20260504.docx")
)

# ── Public entry points ───────────────────────────────────────────────────────

def compile_submittal(body: dict, *, with_location_links: bool = False) -> io.BytesIO:
    """Build a populated .docx and return it as a seeked BytesIO.
    `with_location_links=True` wraps each "Location: <label>" in a hyperlink to
    a marker URI that the PDF post-processor rewrites into an internal /GoTo.
    """
    fields   = body.get("fields",   {})
    comments = body.get("comments", [])
    markup   = body.get("markup",   {})

    with open(TEMPLATE_PATH, "rb") as f:
        template_bytes = f.read()

    doc = Document(io.BytesIO(template_bytes))

    _fill_sdts(doc, fields)
    _set_response_requested(doc, bool(fields.get("response_requested", False)))
    _fill_comment_tables(doc, comments)
    _fill_markup_sections(doc, markup, with_location_links=with_location_links)
    _sync_page_borders(doc)

    modified_doc_xml = etree.tostring(
        doc.element, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # When we add hyperlink relationships via doc.part.relate_to, those new
    # rIds live only in python-docx memory unless we re-serialize the rels
    # file. Without this, document.xml references unknown rIds and Word
    # opens the file as corrupted.
    rels_xml = _serialize_part_rels(doc.part) if with_location_links else None

    return _rebuild_zip(template_bytes, modified_doc_xml, fields, body, rels_xml=rels_xml)


def compile_submittal_pdf(body: dict, original_pdf: bytes) -> bytes:
    """Build the merged PDF: DOCX → PDF + sanitized appendix + hyperlinks."""
    from services.pdf_finalize import sanitize_pdf, docx_to_pdf, merge_and_hyperlink

    docx_buf   = compile_submittal(body, with_location_links=True)
    docx_bytes = docx_buf.getvalue()

    appendix_bytes, page_map = sanitize_pdf(original_pdf)
    front_pdf = docx_to_pdf(docx_bytes)

    targets: list[tuple[str, int]] = []
    markup = body.get("markup", {}) or {}
    for bucket_key in ("rfi", "equal_substitutions", "deviations", "exclusions"):
        for item in markup.get(bucket_key, []) or []:
            label    = item.get("page", "")
            orig_idx = item.get("page_idx")
            if label and orig_idx is not None and orig_idx in page_map:
                targets.append((label, page_map[orig_idx]))

    return merge_and_hyperlink(front_pdf, appendix_bytes, targets)


def load_submittal(docx_bytes: bytes) -> dict | None:
    """Extract LISA form payload from a previously compiled .docx. Returns None if not found."""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        names = z.namelist()
        if "word/lisa_payload.json" in names:
            return json.loads(z.read("word/lisa_payload.json").decode("utf-8"))
        # Fallback: custom document property written for Word-save durability
        if "docProps/custom.xml" in names:
            text = z.read("docProps/custom.xml").decode("utf-8")
            m = re.search(r'name="LISA_Payload"><vt:lpwstr>(.*?)</vt:lpwstr>', text, re.DOTALL)
            if m:
                return json.loads(html.unescape(m.group(1)))
    return None


# ── ZIP reconstruction ────────────────────────────────────────────────────────

def _rebuild_zip(
    template_bytes: bytes,
    modified_doc_xml: bytes,
    fields: dict,
    payload: dict,
    *,
    rels_xml: bytes | None = None,
) -> io.BytesIO:
    """
    Build the output DOCX by copying every part from the template verbatim,
    with these targeted exceptions:
      - word/document.xml             → replaced by python-docx modified bytes
      - word/_rels/document.xml.rels  → replaced by `rels_xml` if provided
      - docProps/custom.xml           → property values updated from form fields
      - word/settings.xml             → w:updateFields added so Word refreshes on open
      - [Content_Types].xml           → lisa_payload.json Override entry inserted
    Then appends word/lisa_payload.json.
    """
    payload_json = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    ct_entry = (
        '<Override PartName="/word/lisa_payload.json"'
        ' ContentType="application/json"/>'
    )

    in_buf  = io.BytesIO(template_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, "r") as zin, \
         zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, modified_doc_xml)

            elif rels_xml is not None and item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)

            elif item.filename == "docProps/custom.xml":
                data = _patch_custom_properties(zin.read(item.filename), fields, payload)
                zout.writestr(item, data)

            elif item.filename == "word/settings.xml":
                data = zin.read(item.filename).decode("utf-8")
                if "w:updateFields" not in data:
                    data = data.replace("</w:settings>",
                                        '<w:updateFields w:val="1"/></w:settings>')
                zout.writestr(item, data.encode("utf-8"))

            elif item.filename == "[Content_Types].xml":
                data = zin.read(item.filename).decode("utf-8")
                if ct_entry not in data:
                    data = data.replace("</Types>", ct_entry + "</Types>")
                zout.writestr(item, data.encode("utf-8"))

            else:
                zout.writestr(item, zin.read(item.filename))

        zout.writestr("word/lisa_payload.json", payload_json)

    out_buf.seek(0)
    return out_buf


# ── Document part relationships ───────────────────────────────────────────────

def _serialize_part_rels(part) -> bytes:
    """Render `part.rels` (which already includes the original template rels
    plus any added via `part.relate_to`) back to the OOXML rels XML format.
    Used to keep document.xml's r:id references and document.xml.rels in sync
    when we add hyperlinks at compile time.
    """
    PKG_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    root   = etree.Element(f"{{{PKG_NS}}}Relationships", nsmap={None: PKG_NS})
    for rel in part.rels.values():
        elem = etree.SubElement(root, f"{{{PKG_NS}}}Relationship")
        elem.set("Id",   rel.rId)
        elem.set("Type", rel.reltype)
        elem.set("Target", rel.target_ref)
        if rel.is_external:
            elem.set("TargetMode", "External")
    return (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        + etree.tostring(root)
    )


# ── Custom document properties ────────────────────────────────────────────────

def _patch_custom_properties(xml_bytes: bytes, fields: dict, payload: dict | None = None) -> bytes:
    """Write form field values into the Word custom document properties.
    Headers and body display these via { DOCPROPERTY "Name" } field codes.
    Also writes LISA_Payload as a backup so Load survives a Word save.
    """
    text = xml_bytes.decode("utf-8")

    prop_map = {
        "Project Name":    fields.get("project_name",    ""),
        "Project Number":  fields.get("project_number",  ""),
        "End Customer":    fields.get("end_customer",    ""),
        "Site Name":       fields.get("site_name",       ""),
        "Release Package": fields.get("release_package", ""),
        "Date":            fields.get("date",            ""),
    }
    for prop_name, value in prop_map.items():
        escaped = html.escape(value, quote=False)
        text = re.sub(
            rf'(name="{re.escape(prop_name)}"><vt:lpwstr>)[^<]*(</vt:lpwstr>)',
            rf'\g<1>{escaped}\g<2>',
            text,
        )

    filename = fields.get("filename", "").strip()
    if filename.lower().endswith(".docx"):
        filename = filename[:-5]
    if filename:
        text = re.sub(
            r'(name="File Name"><vt:lpwstr>)[^<]*(</vt:lpwstr>)',
            rf'\g<1>{html.escape(filename, quote=False)}\g<2>',
            text,
        )

    if payload is not None:
        payload_escaped = html.escape(json.dumps(payload, ensure_ascii=False), quote=False)
        if 'name="LISA_Payload"' in text:
            text = re.sub(
                r'(name="LISA_Payload"><vt:lpwstr>).*?(</vt:lpwstr>)',
                rf'\g<1>{payload_escaped}\g<2>',
                text, flags=re.DOTALL,
            )
        else:
            pids = re.findall(r'pid="(\d+)"', text)
            next_pid = max(int(p) for p in pids) + 1 if pids else 2
            entry = (
                f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}"'
                f' pid="{next_pid}" name="LISA_Payload">'
                f'<vt:lpwstr>{payload_escaped}</vt:lpwstr></property>'
            )
            text = text.replace("</Properties>", entry + "</Properties>")

    return text.encode("utf-8")


# ── SDT (content control) helpers ─────────────────────────────────────────────

def _find_sdts_by_tag(root_el, tag_name: str) -> list:
    return [
        sdt for sdt in root_el.iter(qn("w:sdt"))
        if (lambda t: t is not None and t.get(qn("w:val")) == tag_name)(
            sdt.find(".//" + qn("w:tag"))
        )
    ]


def _set_sdt_text(sdt_el, text: str):
    content = sdt_el.find(qn("w:sdtContent"))
    if content is None:
        return
    pr = sdt_el.find(qn("w:sdtPr"))
    if pr is not None:
        showing = pr.find(qn("w:showingPlcHdr"))
        if showing is not None:
            pr.remove(showing)
    for r_el in content.iter(qn("w:r")):
        rPr = r_el.find(qn("w:rPr"))
        if rPr is not None:
            for rStyle in rPr.findall(qn("w:rStyle")):
                if rStyle.get(qn("w:val")) == "PlaceholderText":
                    rPr.remove(rStyle)
    first = True
    for t_el in content.iter(qn("w:t")):
        t_el.text = text if first else ""
        first = False


def _fill_sdts(doc: Document, fields: dict):
    body = doc.element.body

    for sdt in _find_sdts_by_tag(body, "Submittal No."):
        _set_sdt_text(sdt, fields.get("revision", "00"))

    for sdt in body.iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        if pr is not None and pr.find(qn("w:date")) is not None:
            _set_sdt_text(sdt, fields.get("date", ""))
            break

    for sdt in _find_sdts_by_tag(body, "Subject"):
        _set_sdt_text(sdt, fields.get("subject", ""))


# ── Response Requested YES / NO checkboxes ────────────────────────────────────

def _set_response_requested(doc: Document, yes: bool):
    body = doc.element.body
    checkboxes = []
    for sdt in body.iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        if pr is None:
            continue
        if any(child.tag.endswith("}checkbox") for child in pr):
            checkboxes.append(sdt)

    checked_char   = "☒"
    unchecked_char = "☐"

    for i, sdt in enumerate(checkboxes[:2]):
        is_checked = (i == 0 and yes) or (i == 1 and not yes)
        display    = checked_char if is_checked else unchecked_char

        content = sdt.find(qn("w:sdtContent"))
        if content is not None:
            for t_el in content.iter(qn("w:t")):
                if t_el.text in (checked_char, unchecked_char, "☐", "☒"):
                    t_el.text = display
                    break

        pr = sdt.find(qn("w:sdtPr"))
        if pr is not None:
            for child in pr:
                if child.tag.endswith("}checkbox"):
                    ns = child.tag.split("}")[0].lstrip("{")
                    checked_el = child.find(f"{{{ns}}}checked")
                    if checked_el is None:
                        checked_el = etree.SubElement(child, f"{{{ns}}}checked")
                    checked_el.set(f"{{{ns}}}val", "1" if is_checked else "0")
                    break


# ── Comment tables ───────────────────────────────────────────────────────────

def _fill_comment_tables(doc: Document, comments: list):
    """Fill or clone the template comment table for each comment block."""
    if len(doc.tables) < 3:
        return
    template_table = doc.tables[2]
    pristine_tbl   = copy.deepcopy(template_table._tbl)

    if not comments:
        rows = list(template_table._tbl.findall(qn("w:tr")))
        for row in rows[3:]:
            row.getparent().remove(row)
        return

    _fill_single_comment(template_table._tbl, comments[0], 1)
    prev_tbl = template_table._tbl
    for idx, comment in enumerate(comments[1:], start=2):
        spacer = _make_normal_spacer()
        prev_tbl.addnext(spacer)
        new_tbl = copy.deepcopy(pristine_tbl)
        spacer.addnext(new_tbl)
        _fill_single_comment(new_tbl, comment, idx)
        prev_tbl = new_tbl


def _fill_single_comment(tbl_el, comment: dict, idx: int):
    """Populate one comment w:tbl element (lxml) from a comment dict."""
    rows = list(tbl_el.findall(qn("w:tr")))
    if len(rows) < 3:
        return

    # Row 0:
    #   Cell 1 holds a SEQ auto-number field — leave it; Word renumbers on field update.
    #   Status is in a row-level SDT (tag="Status"), not in cell 2 (which is the label).
    for sdt in rows[0].iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        if pr is not None:
            tag_el = pr.find(qn("w:tag"))
            if tag_el is not None and tag_el.get(qn("w:val")) == "Status":
                _set_sdt_text(sdt, comment.get("status", "OPEN"))
                break

    # Row 1:
    #   Cell 0 is the "Document" label — leave it.
    #   Document name is in a row-level SDT (tag="Document Name").
    for sdt in rows[1].iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        if pr is not None:
            tag_el = pr.find(qn("w:tag"))
            if tag_el is not None and tag_el.get(qn("w:val")) == "Document Name":
                _set_sdt_text(sdt, comment.get("document_name", ""))
                break

    # Row 2: Column headers — leave unchanged

    # Capture pristine data row before removing example rows
    pristine_data_row = copy.deepcopy(rows[3]) if len(rows) >= 4 else None

    # Remove all example data rows (rows 3+)
    for row in rows[3:]:
        row.getparent().remove(row)

    if pristine_data_row is None:
        return

    for entry in comment.get("rows", []):
        new_tr = copy.deepcopy(pristine_data_row)
        cells  = [el for el in new_tr if el.tag == qn("w:tc")]
        if len(cells) >= 3:
            _set_tr_cell_text(cells[0], entry.get("by",       ""))
            _set_tr_cell_text(cells[1], entry.get("revision", ""))
            _set_tr_cell_text(cells[2], entry.get("comment",  ""))
        tbl_el.append(new_tr)


def _set_tr_cell_text(tc_el, text: str):
    """Set the text of a w:tc element obtained directly from lxml.
    Clears all para content except w:pPr to handle fldSimple, w:ins, and other wrappers.
    """
    para = tc_el.find(qn("w:p"))
    if para is None:
        return
    pPr = para.find(qn("w:pPr"))
    for child in list(para):
        para.remove(child)
    if pPr is not None:
        para.insert(0, pPr)
    r_el = etree.SubElement(para, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


# ── Markup Summary sections ───────────────────────────────────────────────────

_MARKUP_SECTIONS = [
    ("Request for Information", "rfi"),
    ("Equal Substitutions",     "equal_substitutions"),
    ("Deviations",              "deviations"),
    ("Exclusions",              "exclusions"),
]

_SECTION_HEADING_TEXTS = {h for h, _ in _MARKUP_SECTIONS}
_STOP_STYLES = {"Heading 1 Plain", "Heading 2 Plain", "Heading 1", "Heading 2"}


def _fill_markup_sections(doc: Document, markup: dict, *, with_location_links: bool = False):
    for heading_text, bucket_key in _MARKUP_SECTIONS:
        items = markup.get(bucket_key, [])
        _rebuild_markup_section(doc, heading_text, items, with_location_links=with_location_links)


def _rebuild_markup_section(doc: Document, heading_text: str, items: list, *, with_location_links: bool = False):
    body         = doc.element.body
    all_children = list(body.iterchildren())  # paragraphs AND tables

    heading_el  = None
    heading_idx = -1
    for idx, el in enumerate(all_children):
        if el.tag != qn("w:p"):
            continue
        texts = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if texts.strip() == heading_text:
            heading_el  = el
            heading_idx = idx
            break

    if heading_el is None:
        return

    end_idx = len(all_children)
    for idx in range(heading_idx + 1, len(all_children)):
        el = all_children[idx]
        if el.tag != qn("w:p"):
            continue
        texts     = "".join(t.text or "" for t in el.iter(qn("w:t")))
        style_val = ""
        pPr = el.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val"), "")
        if texts.strip() in _SECTION_HEADING_TEXTS or style_val in _STOP_STYLES:
            end_idx = idx
            break

    for el in all_children[heading_idx + 1: end_idx]:
        if el.tag == qn("w:sectPr"):
            continue
        if el.getparent() is not None:
            el.getparent().remove(el)

    if not items:
        return

    insert_after = heading_el
    for item_num, item in enumerate(items, 1):
        label = item.get("page", "")
        if with_location_links and label:
            location_para = _make_location_para_with_link(doc, label)
        else:
            location_para = _make_no_spacing_para(f"Location: {label}")
        blocks = [
            _make_no_spacing_para(f"Item: {item_num:03d}"),
            location_para,
            _make_no_spacing_para(f"Comment: {item.get('comment', '')}"),
            _make_normal_para(""),
        ]
        for p_el in blocks:
            insert_after.addnext(p_el)
            insert_after = p_el


# ── Paragraph factory helpers ─────────────────────────────────────────────────

# Marker URI used to round-trip "Location:" hyperlinks through DOCX -> PDF.
# pdf_finalize.merge_and_hyperlink rewrites these to /GoTo internal links.
_LOCATION_LINK_PREFIX = "https://lisa.invalid/page/"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _make_location_para_with_link(doc: Document, label: str):
    """Build a NoSpacing paragraph: 'Location: ' + hyperlinked <label>.
    The hyperlink target is a marker URI that the PDF post-processor will
    rewrite into an internal /GoTo to the appendix page; if no matching
    appendix page exists, the post-processor drops the link annotation
    (the styled text remains).
    """
    from urllib.parse import quote
    from docx.opc.constants import RELATIONSHIP_TYPE

    p_el = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p_el, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "NoSpacing")

    r1 = etree.SubElement(p_el, qn("w:r"))
    t1 = etree.SubElement(r1, qn("w:t"))
    t1.text = "Location: "
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    uri = _LOCATION_LINK_PREFIX + quote(label, safe="")
    rId = doc.part.relate_to(uri, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = etree.SubElement(p_el, qn("w:hyperlink"))
    hyperlink.set(f"{{{_R_NS}}}id", rId)

    r2 = etree.SubElement(hyperlink, qn("w:r"))
    rPr = etree.SubElement(r2, qn("w:rPr"))
    rStyle = etree.SubElement(rPr, qn("w:rStyle"))
    rStyle.set(qn("w:val"), "Hyperlink")
    t2 = etree.SubElement(r2, qn("w:t"))
    t2.text = label

    return p_el


def _make_no_spacing_para(text: str):
    p      = etree.Element(qn("w:p"))
    pPr    = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "NoSpacing")
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def _make_normal_para(text: str):
    p = etree.Element(qn("w:p"))
    if text:
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
    return p


def _make_normal_spacer():
    """Empty paragraph with explicit Normal style — used between cloned comment tables."""
    p     = etree.Element(qn("w:p"))
    pPr   = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "Normal")
    return p


def _sync_page_borders(doc: Document):
    """Copy pgBorders from the intermediate sectPr (section 1) into the final sectPr (section 2).
    Iterates direct body children only to avoid finding sectPrs inside table cells.
    """
    body = doc.element.body
    source_borders = None
    for child in body:
        if child.tag != qn("w:p"):
            continue
        pPr = child.find(qn("w:pPr"))
        if pPr is None:
            continue
        sp = pPr.find(qn("w:sectPr"))
        if sp is None:
            continue
        pb = sp.find(qn("w:pgBorders"))
        if pb is not None:
            source_borders = pb
            break

    if source_borders is None:
        return

    final_sect = body.find(qn("w:sectPr"))
    if final_sect is None:
        return

    existing = final_sect.find(qn("w:pgBorders"))
    if existing is not None:
        final_sect.remove(existing)
    final_sect.append(copy.deepcopy(source_borders))
